"""
Word document automation tools — read and create .docx files via python-docx.

Windows only: uses local filesystem paths. No COM/Outlook dependency.
"""

from __future__ import annotations

import json
import logging
import os
from pathlib import Path

from langchain.tools import tool

logger = logging.getLogger(__name__)


def _resolve_path(filepath: str) -> str:
    """Normalize and validate a file path."""
    path = Path(os.path.normpath(filepath)).resolve()
    return str(path)


def _check_docx_installed() -> str | None:
    """Return an error message if python-docx is not available, else None."""
    try:
        import docx  # noqa: F401
    except ImportError:
        return "python-docx is not installed. Run: pip install python-docx"
    return None


@tool("read_word_document", parse_docstring=True)
def read_word_document_tool(filepath: str) -> str:
    """Read the full text content of a Word (.docx) document.

    Returns each paragraph as a separate element in the JSON response.

    Args:
        filepath: Absolute or relative path to the .docx file to read.
    """
    err = _check_docx_installed()
    if err:
        return json.dumps({"error": err}, ensure_ascii=False)

    path = _resolve_path(filepath)
    if not os.path.isfile(path):
        return json.dumps({"error": f"File not found: {path}"}, ensure_ascii=False)
    if not path.lower().endswith(".docx"):
        return json.dumps({"error": f"Unsupported format. Only .docx files are supported: {path}"}, ensure_ascii=False)

    try:
        import docx

        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        # Also extract tables
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)

        return json.dumps(
            {
                "filepath": path,
                "total_paragraphs": len(paragraphs),
                "paragraphs": paragraphs,
                "tables": tables,
                "total_tables": len(tables),
            },
            indent=2,
            ensure_ascii=False,
        )
    except Exception as exc:
        logger.exception("Failed to read Word document: %s", exc)
        return json.dumps({"error": f"Failed to read Word document: {exc}"}, ensure_ascii=False)


@tool("create_word_document", parse_docstring=True)
def create_word_document_tool(filepath: str, content: str) -> str:
    """Create or overwrite a Word (.docx) document with the given text content.

    Each line of content becomes a paragraph. Supports basic **bold** and
    *italic* markers (converted to actual Word formatting).

    Args:
        filepath: Absolute or relative path where the .docx file will be saved.
        content: Text content. Lines separated by newline become paragraphs.
    """
    err = _check_docx_installed()
    if err:
        return json.dumps({"error": err}, ensure_ascii=False)

    path = _resolve_path(filepath)
    if not path.lower().endswith(".docx"):
        path += ".docx"

    try:
        import re

        import docx
        from docx.shared import Pt

        doc = docx.Document()

        for line in content.split("\n"):
            if not line.strip():
                doc.add_paragraph("")
                continue

            # Parse basic markdown-style formatting
            parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", line)
            para = doc.add_paragraph()
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                else:
                    para.add_run(part)

        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        doc.save(path)

        return json.dumps(
            {
                "filepath": path,
                "status": "created",
                "paragraphs": len(content.split("\n")),
            },
            indent=2,
            ensure_ascii=False,
        )
    except Exception as exc:
        logger.exception("Failed to create Word document: %s", exc)
        return json.dumps({"error": f"Failed to create Word document: {exc}"}, ensure_ascii=False)

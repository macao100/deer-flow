"""
Serveur MCP Office — expose les outils Word, Excel et Outlook via le protocole MCP.

Usage :
  python office_mcp_server.py

Le serveur communique en stdio (JSON-RPC), compatible avec tout client MCP.
Les outils sont les mêmes que ceux utilisés en interne par DeerFlow.
"""

from __future__ import annotations

import json
import logging
import os
from datetime import datetime, timedelta
from pathlib import Path

from mcp.server import Server
from mcp.server.stdio import stdio_server

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _resolve_path(filepath: str) -> str:
    """Normalize and validate a file path."""
    return str(Path(os.path.normpath(filepath)).resolve())


# ---------------------------------------------------------------------------
# Word tools
# ---------------------------------------------------------------------------


def _read_word_document(filepath: str) -> str:
    """Read the full text content of a Word (.docx) document."""
    try:
        import docx
    except ImportError:
        return json.dumps({"error": "python-docx is not installed. Run: pip install python-docx"}, ensure_ascii=False)

    path = _resolve_path(filepath)
    if not os.path.isfile(path):
        return json.dumps({"error": f"File not found: {path}"}, ensure_ascii=False)
    if not path.lower().endswith(".docx"):
        return json.dumps({"error": f"Unsupported format. Only .docx files are supported: {path}"}, ensure_ascii=False)

    try:
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        tables = []
        for table in doc.tables:
            tables.append([[cell.text for cell in row.cells] for row in table.rows])

        return json.dumps(
            {
                "filepath": path,
                "total_paragraphs": len(paragraphs),
                "paragraphs": paragraphs,
                "tables": tables,
                "total_tables": len(tables),
            },
            indent=2,
            ensure_ascii=False,
        )
    except Exception as exc:
        return json.dumps({"error": f"Failed to read Word document: {exc}"}, ensure_ascii=False)


def _create_word_document(filepath: str, content: str) -> str:
    """Create or overwrite a Word (.docx) document."""
    try:
        import docx
    except ImportError:
        return json.dumps({"error": "python-docx is not installed. Run: pip install python-docx"}, ensure_ascii=False)

    path = _resolve_path(filepath)
    if not path.lower().endswith(".docx"):
        path += ".docx"

    try:
        import re

        doc = docx.Document()
        for line in content.split("\n"):
            if not line.strip():
                doc.add_paragraph("")
                continue
            parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", line)
            para = doc.add_paragraph()
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                else:
                    para.add_run(part)

        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        doc.save(path)

        return json.dumps(
            {"filepath": path, "status": "created", "paragraphs": len(content.split("\n"))},
            indent=2,
            ensure_ascii=False,
        )
    except Exception as exc:
        return json.dumps({"error": f"Failed to create Word document: {exc}"}, ensure_ascii=False)


# ---------------------------------------------------------------------------
# Excel tools
# ---------------------------------------------------------------------------


def _read_excel_range(filepath: str, sheet_name: str | None = None) -> str:
    """Read data from an Excel (.xlsx) workbook."""
    try:
        import openpyxl
    except ImportError:
        return json.dumps({"error": "openpyxl is not installed. Run: pip install openpyxl"}, ensure_ascii=False)

    path = _resolve_path(filepath)
    if not os.path.isfile(path):
        return json.dumps({"error": f"File not found: {path}"}, ensure_ascii=False)

    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        if sheet_name:
            if sheet_name not in wb.sheetnames:
                wb.close()
                return json.dumps({"error": f"Sheet '{sheet_name}' not found. Available: {wb.sheetnames}"}, ensure_ascii=False)
            ws = wb[sheet_name]
        else:
            ws = wb.active
            sheet_name = ws.title

        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([str(cell) if cell is not None else None for cell in row])
        wb.close()

        return json.dumps(
            {
                "filepath": path,
                "sheet_name": sheet_name,
                "total_rows": len(rows),
                "total_columns": len(rows[0]) if rows else 0,
                "data": rows,
            },
            indent=2,
            ensure_ascii=False,
        )
    except Exception as exc:
        return json.dumps({"error": f"Failed to read Excel file: {exc}"}, ensure_ascii=False)


def _write_excel_range(filepath: str, sheet_name: str | None = None, data: str = "[]") -> str:
    """Write data to an Excel (.xlsx) file."""
    try:
        import openpyxl
    except ImportError:
        return json.dumps({"error": "openpyxl is not installed. Run: pip install openpyxl"}, ensure_ascii=False)

    path = _resolve_path(filepath)
    if not path.lower().endswith(".xlsx"):
        path += ".xlsx"

    try:
        try:
            parsed = json.loads(data)
            if not isinstance(parsed, list):
                return json.dumps({"error": "data must be a JSON array of rows (list of lists)"}, ensure_ascii=False)
        except json.JSONDecodeError as e:
            return json.dumps({"error": f"Invalid JSON data: {e}"}, ensure_ascii=False)

        if os.path.isfile(path):
            wb = openpyxl.load_workbook(path)
        else:
            wb = openpyxl.Workbook()
            wb.remove(wb.active)

        if sheet_name and sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    cell.value = None
        else:
            ws = wb.create_sheet(title=sheet_name) if sheet_name else wb.create_sheet(title="Sheet1")

        for row_idx, row_data in enumerate(parsed, start=1):
            if isinstance(row_data, list):
                for col_idx, value in enumerate(row_data, start=1):
                    if value is not None:
                        ws.cell(row=row_idx, column=col_idx).value = value

        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        wb.save(path)
        wb.close()

        return json.dumps(
            {"filepath": path, "sheet_name": ws.title, "rows_written": len(parsed), "status": "created" if not os.path.isfile(path) else "updated"},
            indent=2,
            ensure_ascii=False,
        )
    except Exception as exc:
        return json.dumps({"error": f"Failed to write Excel file: {exc}"}, ensure_ascii=False)


def _list_excel_sheets(filepath: str) -> str:
    """List all sheet names in an Excel (.xlsx) workbook."""
    try:
        import openpyxl
    except ImportError:
        return json.dumps({"error": "openpyxl is not installed. Run: pip install openpyxl"}, ensure_ascii=False)

    path = _resolve_path(filepath)
    if not os.path.isfile(path):
        return json.dumps({"error": f"File not found: {path}"}, ensure_ascii=False)

    try:
        wb = openpyxl.load_workbook(path, read_only=True)
        result = {
            "filepath": path,
            "sheets": wb.sheetnames,
            "total_sheets": len(wb.sheetnames),
        }
        wb.close()
        return json.dumps(result, indent=2, ensure_ascii=False)
    except Exception as exc:
        return json.dumps({"error": f"Failed to list Excel sheets: {exc}"}, ensure_ascii=False)


# ---------------------------------------------------------------------------
# Outlook tools
# ---------------------------------------------------------------------------


def _get_outlook():
    """Get the Outlook Application object via COM."""
    import win32com.client
    return win32com.client.Dispatch("Outlook.Application")


def _send_outlook_email(to: str, subject: str = "", body: str = "", cc: str = "", bcc: str = "", attachment_paths: str = "") -> str:
    """Send an email through Microsoft Outlook (displays for review)."""
    try:
        import win32com.client  # noqa: F401
    except ImportError:
        return json.dumps({"error": "pywin32 is not installed. Run: pip install pywin32"}, ensure_ascii=False)

    if not to or not to.strip():
        return json.dumps({"error": "Recipient (to) is required"}, ensure_ascii=False)

    try:
        outlook = _get_outlook()
        mail = outlook.CreateItem(0)
        mail.To = to.strip()
        if cc.strip():
            mail.CC = cc.strip()
        if bcc.strip():
            mail.BCC = bcc.strip()
        mail.Subject = subject.strip() if subject.strip() else "(No subject)"
        mail.Body = body.strip() if body.strip() else ""

        if attachment_paths.strip():
            for p in attachment_paths.split(";"):
                p = p.strip()
                if p and os.path.isfile(p):
                    mail.Attachments.Add(os.path.abspath(p))

        mail.Display(True)

        return json.dumps(
            {
                "status": "displayed_for_review",
                "to": mail.To,
                "subject": mail.Subject,
                "message": "Email opened in Outlook. Review and click Send to deliver.",
            },
            indent=2,
            ensure_ascii=False,
        )
    except Exception as exc:
        return json.dumps({"error": f"Failed to send email: {exc}"}, ensure_ascii=False)


def _search_outlook_inbox(folder_name: str = "Inbox", max_results: int = 10, unread_only: bool = False, days_back: int = 7, search_term: str = "") -> str:
    """Search emails in an Outlook folder."""
    try:
        import win32com.client  # noqa: F401
    except ImportError:
        return json.dumps({"error": "pywin32 is not installed. Run: pip install pywin32"}, ensure_ascii=False)

    max_results = min(max(max_results, 1), 50)

    try:
        outlook = _get_outlook()
        namespace = outlook.GetNamespace("MAPI")

        ol_default_folders = {
            "Inbox": 6, "Calendar": 9, "Contacts": 10, "Deleted Items": 3,
            "Drafts": 16, "Junk Email": 23, "Outbox": 4, "Sent Items": 5, "Tasks": 13,
        }
        folder = None
        try:
            folder = namespace.Folders.Item(1).Folders(folder_name)
        except Exception:
            ol_type = ol_default_folders.get(folder_name)
            if ol_type is not None:
                folder = namespace.GetDefaultFolder(ol_type)

        if folder is None:
            try:
                available = [f.Name for f in namespace.Folders.Item(1).Folders]
            except Exception:
                available = []
            return json.dumps({"error": f"Folder '{folder_name}' not found.", "available_folders": available}, ensure_ascii=False)

        cutoff = datetime.now() - timedelta(days=days_back)
        items = folder.Items
        items.Sort("[ReceivedTime]", True)

        filters = [f"[ReceivedTime] >= '{cutoff.strftime('%Y-%m-%d')}'"]
        if unread_only:
            filters.append("[UnRead] = True")
        if search_term:
            filters.append(f"@SQL=urn:schemas:httpmail:subject\" ci_startswith '{search_term}'")

        filtered = items.Restrict(" AND ".join(f"({f})" for f in filters)) if len(filters) > 1 else items.Restrict(filters[0])

        emails = []
        for i, item in enumerate(filtered):
            if i >= max_results:
                break
            try:
                received = item.ReceivedTime.strftime("%Y-%m-%d %H:%M") if item.ReceivedTime else ""
            except Exception:
                received = ""
            emails.append({
                "subject": getattr(item, "Subject", "(No subject)"),
                "sender": getattr(item, "SenderName", getattr(item, "SenderEmailAddress", "")),
                "received": received,
                "unread": bool(getattr(item, "UnRead", False)),
                "body_preview": (getattr(item, "Body", "") or "")[:200],
            })

        return json.dumps(
            {"folder": folder_name, "total_found": len(emails), "emails": emails},
            indent=2,
            ensure_ascii=False,
        )
    except Exception as exc:
        return json.dumps({"error": f"Failed to search Outlook: {exc}"}, ensure_ascii=False)


def _list_outlook_folders() -> str:
    """List all available email folders in Outlook."""
    try:
        import win32com.client  # noqa: F401
    except ImportError:
        return json.dumps({"error": "pywin32 is not installed. Run: pip install pywin32"}, ensure_ascii=False)

    try:
        outlook = _get_outlook()
        namespace = outlook.GetNamespace("MAPI")
        store = namespace.Folders.Item(1)
        folders = [{"name": f.Name, "item_count": getattr(f.Items, "Count", 0)} for f in store.Folders]
        return json.dumps({"store_name": store.Name, "folders": folders, "total_folders": len(folders)}, indent=2, ensure_ascii=False)
    except Exception as exc:
        return json.dumps({"error": f"Failed to list Outlook folders: {exc}"}, ensure_ascii=False)


# ---------------------------------------------------------------------------
# MCP Server (low-level API v1)
# ---------------------------------------------------------------------------

import mcp.types as types
from mcp.server.lowlevel import Server

server = Server("office-mcp")

# --- Tool registry ---

TOOLS = {
    "read_word_document": {
        "fn": _read_word_document,
        "description": "Read the full text content of a Word (.docx) document. Returns paragraphs and tables as JSON.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "filepath": {"type": "string", "description": "Absolute or relative path to the .docx file to read."},
            },
            "required": ["filepath"],
        },
    },
    "create_word_document": {
        "fn": _create_word_document,
        "description": "Create or overwrite a Word (.docx) document. Each line becomes a paragraph. Supports **bold** and *italic* markers.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "filepath": {"type": "string", "description": "Path where the .docx file will be saved."},
                "content": {"type": "string", "description": "Text content. Lines separated by newline become paragraphs."},
            },
            "required": ["filepath", "content"],
        },
    },
    "read_excel_range": {
        "fn": _read_excel_range,
        "description": "Read data from an Excel (.xlsx) workbook. Returns all rows as a JSON array.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "filepath": {"type": "string", "description": "Path to the .xlsx file."},
                "sheet_name": {"type": "string", "description": "Sheet name (optional, reads first sheet by default)."},
            },
            "required": ["filepath"],
        },
    },
    "write_excel_range": {
        "fn": _write_excel_range,
        "description": "Write data to an Excel (.xlsx) file. Creates file if it doesn't exist.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "filepath": {"type": "string", "description": "Path where the .xlsx file will be saved."},
                "sheet_name": {"type": "string", "description": "Sheet name (optional)."},
                "data": {"type": "string", "description": "JSON string representing a 2D array (list of rows)."},
            },
            "required": ["filepath"],
        },
    },
    "list_excel_sheets": {
        "fn": _list_excel_sheets,
        "description": "List all sheet names in an Excel (.xlsx) workbook.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "filepath": {"type": "string", "description": "Path to the .xlsx file."},
            },
            "required": ["filepath"],
        },
    },
    "send_outlook_email": {
        "fn": _send_outlook_email,
        "description": "Send an email through Microsoft Outlook. Displays for review (user clicks Send).",
        "inputSchema": {
            "type": "object",
            "properties": {
                "to": {"type": "string", "description": "Recipient email(s), semicolon-separated."},
                "subject": {"type": "string", "description": "Email subject line."},
                "body": {"type": "string", "description": "Email body (plain text)."},
                "cc": {"type": "string", "description": "CC recipient(s), semicolon-separated."},
                "bcc": {"type": "string", "description": "BCC recipient(s), semicolon-separated."},
                "attachment_paths": {"type": "string", "description": "File paths to attach, semicolon-separated."},
            },
            "required": ["to"],
        },
    },
    "search_outlook_inbox": {
        "fn": _search_outlook_inbox,
        "description": "Search emails in an Outlook folder. Returns metadata (sender, subject, date, preview).",
        "inputSchema": {
            "type": "object",
            "properties": {
                "folder_name": {"type": "string", "description": "Outlook folder name (default: 'Inbox')."},
                "max_results": {"type": "integer", "description": "Max emails to return (1-50, default 10)."},
                "unread_only": {"type": "boolean", "description": "Only unread emails if True."},
                "days_back": {"type": "integer", "description": "Days back to search (default 7)."},
                "search_term": {"type": "string", "description": "Optional text filter on subject."},
            },
            "required": [],
        },
    },
    "list_outlook_folders": {
        "fn": _list_outlook_folders,
        "description": "List all available email folders in the default Outlook store.",
        "inputSchema": {
            "type": "object",
            "properties": {},
        },
    },
}


@server.list_tools()
async def handle_list_tools() -> list[types.Tool]:
    """Return the catalog of Office tools."""
    return [
        types.Tool(
            name=name,
            description=info["description"],
            inputSchema=info["inputSchema"],
        )
        for name, info in TOOLS.items()
    ]


@server.call_tool()
async def handle_call_tool(name: str, arguments: dict) -> list[types.TextContent]:
    """Dispatch tool calls to the appropriate implementation."""
    if name not in TOOLS:
        return [types.TextContent(type="text", text=json.dumps({"error": f"Unknown tool: {name}"}, ensure_ascii=False))]

    info = TOOLS[name]
    fn = info["fn"]

    try:
        # Call the tool function with the provided arguments
        result = fn(**arguments)
    except TypeError as exc:
        return [types.TextContent(type="text", text=json.dumps({"error": f"Invalid arguments for {name}: {exc}"}, ensure_ascii=False))]
    except Exception as exc:
        logger.exception("Tool %s failed", name)
        return [types.TextContent(type="text", text=json.dumps({"error": f"Tool {name} failed: {exc}"}, ensure_ascii=False))]

    return [types.TextContent(type="text", text=result)]


def main():
    """Run the MCP server (stdio transport)."""
    import asyncio
    import sys

    logging.basicConfig(level=logging.WARNING, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s", stream=sys.stderr)

    async def run():
        async with stdio_server() as (read, write):
            await server.run(read, write, server.create_initialization_options())

    asyncio.run(run())


if __name__ == "__main__":
    main()
